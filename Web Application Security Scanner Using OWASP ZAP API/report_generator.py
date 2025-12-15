"""
VIP Report Generator
Multiple formats: PDF, HTML, JSON, CSV, DOCX, Excel
Professional 3D styled templates
"""

from datetime import datetime
import json
import csv
import sqlite3
from pathlib import Path

# PDF Generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[!] ReportLab not installed. PDF export disabled.")
    print("    Install with: pip install reportlab")

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[!] python-docx not installed. DOCX export disabled.")
    print("    Install with: pip install python-docx")

# Excel Generation
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("[!] openpyxl not installed. Excel export disabled.")
    print("    Install with: pip install openpyxl")


class VIPReportGenerator:
    def __init__(self, db_path='scan_results.db'):
        self.db_path = db_path
        
    def get_scan_data(self, scan_id):
        """Retrieve scan data from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Get scan info
        cursor.execute('SELECT * FROM scans WHERE id = ?', (scan_id,))
        scan = cursor.fetchone()
        
        if not scan:
            conn.close()
            return None
            
        # Get vulnerabilities
        cursor.execute('SELECT * FROM vulnerabilities WHERE scan_id = ?', (scan_id,))
        vulns = cursor.fetchall()
        
        data = {
            'scan_id': scan[0],
            'target_url': scan[1],
            'scan_type': scan[2],
            'start_time': scan[3],
            'end_time': scan[4],
            'total_alerts': scan[5],
            'high_risk': scan[6],
            'medium_risk': scan[7],
            'low_risk': scan[8],
            'status': scan[9],
            'vulnerabilities': []
        }
        
        for v in vulns:
            data['vulnerabilities'].append({
                'id': v[0],
                'name': v[2],
                'severity': v[3],
                'confidence': v[4],
                'url': v[5],
                'description': v[6],
                'solution': v[7],
                'reference': v[8]
            })
        
        conn.close()
        return data
    
    def generate_html_report(self, scan_id, output_file='report.html'):
        """Generate VIP HTML Report with 3D styling"""
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        html = f'''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Security Scan Report - {data['target_url']}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 40px 20px;
            min-height: 100vh;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }}
        
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 60px 40px;
            text-align: center;
            position: relative;
            overflow: hidden;
        }}
        
        .header::before {{
            content: '';
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
            animation: pulse 15s ease-in-out infinite;
        }}
        
        @keyframes pulse {{
            0%, 100% {{ transform: scale(1); }}
            50% {{ transform: scale(1.1); }}
        }}
        
        .header h1 {{
            font-size: 3em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
            position: relative;
            z-index: 1;
        }}
        
        .header .subtitle {{
            font-size: 1.2em;
            opacity: 0.9;
            position: relative;
            z-index: 1;
        }}
        
        .summary {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 30px;
            padding: 40px;
            background: #f8f9fa;
        }}
        
        .stat-card {{
            background: white;
            padding: 30px;
            border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
            transform: translateY(0);
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        }}
        
        .stat-card::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 5px;
        }}
        
        .stat-card.total::before {{ background: linear-gradient(90deg, #667eea, #764ba2); }}
        .stat-card.high::before {{ background: linear-gradient(90deg, #f093fb, #f5576c); }}
        .stat-card.medium::before {{ background: linear-gradient(90deg, #ffecd2, #fcb69f); }}
        .stat-card.low::before {{ background: linear-gradient(90deg, #a8edea, #fed6e3); }}
        
        .stat-card:hover {{
            transform: translateY(-10px);
            box-shadow: 0 15px 40px rgba(0,0,0,0.2);
        }}
        
        .stat-card h3 {{
            font-size: 0.9em;
            color: #666;
            margin-bottom: 10px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        
        .stat-card .number {{
            font-size: 3em;
            font-weight: bold;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        
        .stat-card.high .number {{
            background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        
        .stat-card.medium .number {{
            background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        
        .stat-card.low .number {{
            background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        
        .info-section {{
            padding: 40px;
            background: white;
        }}
        
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        
        .info-item {{
            padding: 20px;
            background: #f8f9fa;
            border-radius: 10px;
            border-left: 4px solid #667eea;
        }}
        
        .info-item label {{
            font-weight: bold;
            color: #667eea;
            display: block;
            margin-bottom: 5px;
        }}
        
        .info-item value {{
            color: #333;
        }}
        
        .vulnerabilities {{
            padding: 40px;
            background: white;
        }}
        
        .section-title {{
            font-size: 2em;
            margin-bottom: 30px;
            color: #333;
            text-align: center;
            position: relative;
            padding-bottom: 15px;
        }}
        
        .section-title::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 50%;
            transform: translateX(-50%);
            width: 100px;
            height: 4px;
            background: linear-gradient(90deg, #667eea, #764ba2);
            border-radius: 2px;
        }}
        
        .vuln-card {{
            background: white;
            border-radius: 15px;
            padding: 30px;
            margin-bottom: 30px;
            box-shadow: 0 5px 20px rgba(0,0,0,0.1);
            border-left: 6px solid #ddd;
            transition: all 0.3s ease;
        }}
        
        .vuln-card:hover {{
            box-shadow: 0 10px 30px rgba(0,0,0,0.15);
            transform: translateX(5px);
        }}
        
        .vuln-card.high {{ border-left-color: #f5576c; }}
        .vuln-card.medium {{ border-left-color: #fcb69f; }}
        .vuln-card.low {{ border-left-color: #a8edea; }}
        
        .vuln-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
        }}
        
        .vuln-title {{
            font-size: 1.5em;
            color: #333;
            font-weight: bold;
        }}
        
        .severity-badge {{
            padding: 8px 20px;
            border-radius: 25px;
            font-weight: bold;
            font-size: 0.9em;
            text-transform: uppercase;
            letter-spacing: 1px;
            box-shadow: 0 4px 10px rgba(0,0,0,0.2);
        }}
        
        .severity-badge.high {{
            background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
            color: white;
        }}
        
        .severity-badge.medium {{
            background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
            color: #8B4513;
        }}
        
        .severity-badge.low {{
            background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
            color: #333;
        }}
        
        .vuln-content {{
            color: #666;
            line-height: 1.8;
        }}
        
        .vuln-content p {{
            margin-bottom: 15px;
        }}
        
        .vuln-content strong {{
            color: #333;
            display: inline-block;
            margin-right: 10px;
        }}
        
        .solution-box {{
            background: linear-gradient(135deg, #d4fc79 0%, #96e6a1 100%);
            padding: 20px;
            border-radius: 10px;
            margin-top: 20px;
            border-left: 4px solid #4CAF50;
        }}
        
        .solution-box strong {{
            color: #2e7d32;
            font-size: 1.1em;
        }}
        
        .footer {{
            background: #2c3e50;
            color: white;
            padding: 40px;
            text-align: center;
        }}
        
        .footer p {{
            margin-bottom: 10px;
            opacity: 0.8;
        }}
        
        .btn-3d {{
            display: inline-block;
            padding: 15px 40px;
            margin: 10px;
            border-radius: 50px;
            font-weight: bold;
            text-decoration: none;
            color: white;
            box-shadow: 0 8px 15px rgba(0,0,0,0.3);
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        }}
        
        .btn-3d::before {{
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: rgba(255,255,255,0.2);
            transition: all 0.3s ease;
        }}
        
        .btn-3d:hover::before {{
            left: 100%;
        }}
        
        .btn-3d:hover {{
            transform: translateY(-5px);
            box-shadow: 0 15px 30px rgba(0,0,0,0.4);
        }}
        
        .btn-3d:active {{
            transform: translateY(-2px);
            box-shadow: 0 5px 10px rgba(0,0,0,0.3);
        }}
        
        .btn-primary {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        }}
        
        .btn-success {{
            background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
        }}
        
        .btn-danger {{
            background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        }}
        
        .action-buttons {{
            text-align: center;
            padding: 40px;
            background: #f8f9fa;
        }}
        
        @media print {{
            .action-buttons, .btn-3d {{ display: none; }}
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            <h1>🛡️ SECURITY SCAN REPORT</h1>
            <p class="subtitle">Comprehensive Vulnerability Assessment</p>
        </div>
        
        <!-- Summary Cards -->
        <div class="summary">
            <div class="stat-card total">
                <h3>Total Issues</h3>
                <div class="number">{data['total_alerts']}</div>
            </div>
            <div class="stat-card high">
                <h3>High Risk</h3>
                <div class="number">{data['high_risk']}</div>
            </div>
            <div class="stat-card medium">
                <h3>Medium Risk</h3>
                <div class="number">{data['medium_risk']}</div>
            </div>
            <div class="stat-card low">
                <h3>Low Risk</h3>
                <div class="number">{data['low_risk']}</div>
            </div>
        </div>
        
        <!-- Scan Information -->
        <div class="info-section">
            <h2 class="section-title">Scan Information</h2>
            <div class="info-grid">
                <div class="info-item">
                    <label>Target URL:</label>
                    <value>{data['target_url']}</value>
                </div>
                <div class="info-item">
                    <label>Scan Type:</label>
                    <value>{data['scan_type'].title()}</value>
                </div>
                <div class="info-item">
                    <label>Start Time:</label>
                    <value>{data['start_time']}</value>
                </div>
                <div class="info-item">
                    <label>End Time:</label>
                    <value>{data['end_time']}</value>
                </div>
                <div class="info-item">
                    <label>Status:</label>
                    <value>{data['status'].title()}</value>
                </div>
                <div class="info-item">
                    <label>Report Generated:</label>
                    <value>{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</value>
                </div>
            </div>
        </div>
        
        <!-- Vulnerabilities -->
        <div class="vulnerabilities">
            <h2 class="section-title">Detailed Findings</h2>
'''
        
        # Add vulnerabilities
        for idx, vuln in enumerate(data['vulnerabilities'], 1):
            severity_class = vuln['severity'].lower()
            html += f'''
            <div class="vuln-card {severity_class}">
                <div class="vuln-header">
                    <div class="vuln-title">{idx}. {vuln['name']}</div>
                    <div class="severity-badge {severity_class}">{vuln['severity']}</div>
                </div>
                <div class="vuln-content">
                    <p><strong>🔍 Description:</strong> {vuln['description']}</p>
                    <p><strong>📍 Location:</strong> {vuln['url']}</p>
                    <p><strong>🎯 Confidence:</strong> {vuln['confidence']}</p>
'''
            
            if vuln['solution']:
                html += f'''
                    <div class="solution-box">
                        <p><strong>💡 Recommended Solution:</strong></p>
                        <p>{vuln['solution']}</p>
                    </div>
'''
            
            if vuln['reference']:
                html += f'''
                    <p><strong>📚 Reference:</strong> {vuln['reference']}</p>
'''
            
            html += '''
                </div>
            </div>
'''
        
        html += '''
        </div>
        
        <!-- Action Buttons -->
        <div class="action-buttons">
            <a href="#" onclick="window.print(); return false;" class="btn-3d btn-primary">🖨️ Print Report</a>
            <a href="#" onclick="window.location.reload();" class="btn-3d btn-success">🔄 Refresh</a>
            <a href="#" class="btn-3d btn-danger">📧 Email Report</a>
        </div>
        
        <!-- Footer -->
        <div class="footer">
            <p><strong>Generated by Web Security Scanner v1.0</strong></p>
            <p>Powered by OWASP ZAP | Report ID: ''' + str(scan_id) + '''</p>
            <p>© ''' + str(datetime.now().year) + ''' - All Rights Reserved</p>
        </div>
    </div>
    
    <script>
        // Add smooth scroll
        document.querySelectorAll('a[href^="#"]').forEach(anchor => {
            anchor.addEventListener('click', function (e) {
                e.preventDefault();
                document.querySelector(this.getAttribute('href')).scrollIntoView({
                    behavior: 'smooth'
                });
            });
        });
        
        // Add animation on scroll
        const observerOptions = {
            threshold: 0.1,
            rootMargin: '0px 0px -100px 0px'
        };
        
        const observer = new IntersectionObserver((entries) => {
            entries.forEach(entry => {
                if (entry.isIntersecting) {
                    entry.target.style.opacity = '1';
                    entry.target.style.transform = 'translateY(0)';
                }
            });
        }, observerOptions);
        
        document.querySelectorAll('.vuln-card').forEach(card => {
            card.style.opacity = '0';
            card.style.transform = 'translateY(30px)';
            card.style.transition = 'all 0.6s ease';
            observer.observe(card);
        });
    </script>
</body>
</html>
'''
        
        # Save HTML file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html)
        
        print(f"[+] VIP HTML Report generated: {output_file}")
        return True
    
    def generate_pdf_report(self, scan_id, output_file='report.pdf'):
        """Generate PDF Report"""
        if not PDF_AVAILABLE:
            print("[!] PDF generation not available. Install: pip install reportlab")
            return False
        
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        doc = SimpleDocTemplate(output_file, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#764ba2'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph("🛡️ SECURITY SCAN REPORT", title_style))
        story.append(Spacer(1, 20))
        
        # Summary Table
        summary_data = [
            ['Metric', 'Count'],
            ['Total Issues', str(data['total_alerts'])],
            ['High Risk', str(data['high_risk'])],
            ['Medium Risk', str(data['medium_risk'])],
            ['Low Risk', str(data['low_risk'])]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 30))
        
        # Scan Info
        story.append(Paragraph("Scan Information", heading_style))
        info_data = [
            ['Target URL:', data['target_url']],
            ['Scan Type:', data['scan_type'].title()],
            ['Start Time:', data['start_time']],
            ['End Time:', data['end_time']],
            ['Status:', data['status'].title()]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold')
        ]))
        
        story.append(info_table)
        story.append(PageBreak())
        
        # Vulnerabilities
        story.append(Paragraph("Detailed Findings", heading_style))
        story.append(Spacer(1, 20))
        
        for idx, vuln in enumerate(data['vulnerabilities'], 1):
            story.append(Paragraph(f"<b>{idx}. {vuln['name']}</b> [{vuln['severity']}]", styles['Heading3']))
            story.append(Paragraph(f"<b>Description:</b> {vuln['description']}", styles['Normal']))
            story.append(Paragraph(f"<b>Location:</b> {vuln['url']}", styles['Normal']))
            if vuln['solution']:
                story.append(Paragraph(f"<b>Solution:</b> {vuln['solution']}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        print(f"[+] PDF Report generated: {output_file}")
        return True
    
    def generate_json_report(self, scan_id, output_file='report.json'):
        """Generate JSON Report"""
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        print(f"[+] JSON Report generated: {output_file}")
        return True
    
    def generate_csv_report(self, scan_id, output_file='report.csv'):
        """Generate CSV Report"""
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            # Headers
            writer.writerow(['Vulnerability Name', 'Severity', 'Confidence', 'URL', 'Description', 'Solution'])
            
            # Data
            for vuln in data['vulnerabilities']:
                writer.writerow([
                    vuln['name'],
                    vuln['severity'],
                    vuln['confidence'],
                    vuln['url'],
                    vuln['description'],
                    vuln['solution']
                ])
        
        print(f"[+] CSV Report generated: {output_file}")
        return True
    
    def generate_docx_report(self, scan_id, output_file='report.docx'):
        """Generate DOCX Report"""
        if not DOCX_AVAILABLE:
            print("[!] DOCX generation not available. Install: pip install python-docx")
            return False
        
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        doc = Document()
        
        # Title
        title = doc.add_heading('SECURITY SCAN REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Summary', 1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Total Issues'
        table.cell(0, 1).text = str(data['total_alerts'])
        table.cell(1, 0).text = 'High Risk'
        table.cell(1, 1).text = str(data['high_risk'])
        table.cell(2, 0).text = 'Medium Risk'
        table.cell(2, 1).text = str(data['medium_risk'])
        table.cell(3, 0).text = 'Low Risk'
        table.cell(3, 1).text = str(data['low_risk'])
        table.cell(4, 0).text = 'Target URL'
        table.cell(4, 1).text = data['target_url']
        
        # Vulnerabilities
        doc.add_page_break()
        doc.add_heading('Detailed Findings', 1)
        
        for idx, vuln in enumerate(data['vulnerabilities'], 1):
            doc.add_heading(f"{idx}. {vuln['name']}", 2)
            doc.add_paragraph(f"Severity: {vuln['severity']}")
            doc.add_paragraph(f"Description: {vuln['description']}")
            doc.add_paragraph(f"Location: {vuln['url']}")
            if vuln['solution']:
                doc.add_paragraph(f"Solution: {vuln['solution']}")
            doc.add_paragraph('')
        
        doc.save(output_file)
        print(f"[+] DOCX Report generated: {output_file}")
        return True
    
    def generate_excel_report(self, scan_id, output_file='report.xlsx'):
        """Generate Excel Report"""
        if not EXCEL_AVAILABLE:
            print("[!] Excel generation not available. Install: pip install openpyxl")
            return False
        
        data = self.get_scan_data(scan_id)
        if not data:
            print(f"[!] Scan {scan_id} not found")
            return False
        
        wb = openpyxl.Workbook()
        
        # Summary Sheet
        ws_summary = wb.active
        ws_summary.title = 'Summary'
        
        ws_summary['A1'] = 'SECURITY SCAN REPORT'
        ws_summary['A1'].font = Font(size=16, bold=True)
        
        ws_summary['A3'] = 'Target URL'
        ws_summary['B3'] = data['target_url']
        ws_summary['A4'] = 'Scan Type'
        ws_summary['B4'] = data['scan_type']
        ws_summary['A5'] = 'Total Issues'
        ws_summary['B5'] = data['total_alerts']
        ws_summary['A6'] = 'High Risk'
        ws_summary['B6'] = data['high_risk']
        ws_summary['A7'] = 'Medium Risk'
        ws_summary['B7'] = data['medium_risk']
        ws_summary['A8'] = 'Low Risk'
        ws_summary['B8'] = data['low_risk']
        
        # Vulnerabilities Sheet
        ws_vulns = wb.create_sheet('Vulnerabilities')
        headers = ['#', 'Name', 'Severity', 'Confidence', 'URL', 'Description', 'Solution']
        ws_vulns.append(headers)
        
        # Style headers
        for cell in ws_vulns[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='667eea', end_color='667eea', fill_type='solid')
        
        # Add vulnerabilities
        for idx, vuln in enumerate(data['vulnerabilities'], 1):
            ws_vulns.append([
                idx,
                vuln['name'],
                vuln['severity'],
                vuln['confidence'],
                vuln['url'],
                vuln['description'],
                vuln['solution']
            ])
        
        wb.save(output_file)
        print(f"[+] Excel Report generated: {output_file}")
        return True
    
    def generate_all_formats(self, scan_id, base_name='security_report'):
        """Generate reports in all available formats"""
        print("\n" + "="*60)
        print("GENERATING REPORTS IN ALL FORMATS")
        print("="*60)
        
        formats = {
            'HTML': (self.generate_html_report, f'{base_name}.html'),
            'JSON': (self.generate_json_report, f'{base_name}.json'),
            'CSV': (self.generate_csv_report, f'{base_name}.csv'),
        }
        
        if PDF_AVAILABLE:
            formats['PDF'] = (self.generate_pdf_report, f'{base_name}.pdf')
        
        if DOCX_AVAILABLE:
            formats['DOCX'] = (self.generate_docx_report, f'{base_name}.docx')
        
        if EXCEL_AVAILABLE:
            formats['Excel'] = (self.generate_excel_report, f'{base_name}.xlsx')
        
        results = {}
        for format_name, (func, filename) in formats.items():
            print(f"\n[*] Generating {format_name} report...")
            success = func(scan_id, filename)
            results[format_name] = success
        
        print("\n" + "="*60)
        print("REPORT GENERATION SUMMARY")
        print("="*60)
        for format_name, success in results.items():
            status = "✅ SUCCESS" if success else "❌ FAILED"
            print(f"{format_name:12s} : {status}")
        print("="*60 + "\n")
        
        return results


# Main execution
if __name__ == "__main__":
    print("""
    ╔═══════════════════════════════════════════════════════════╗
    ║           VIP REPORT GENERATOR v2.0                      ║
    ║    Multiple Formats: HTML, PDF, JSON, CSV, DOCX, Excel  ║
    ╚═══════════════════════════════════════════════════════════╝
    """)
    
    generator = VIPReportGenerator()
    
    # Get scan ID
    scan_id = input("\nEnter Scan ID: ").strip()
    
    if not scan_id.isdigit():
        print("[!] Invalid Scan ID")
        exit(1)
    
    scan_id = int(scan_id)
    
    print("\nSelect Report Format:")
    print("1. HTML (VIP 3D Design)")
    print("2. PDF")
    print("3. JSON")
    print("4. CSV")
    print("5. DOCX (Word)")
    print("6. Excel")
    print("7. ALL FORMATS")
    
    choice = input("\nYour choice (1-7): ").strip()
    
    if choice == '1':
        generator.generate_html_report(scan_id)
    elif choice == '2':
        generator.generate_pdf_report(scan_id)
    elif choice == '3':
        generator.generate_json_report(scan_id)
    elif choice == '4':
        generator.generate_csv_report(scan_id)
    elif choice == '5':
        generator.generate_docx_report(scan_id)
    elif choice == '6':
        generator.generate_excel_report(scan_id)
    elif choice == '7':
        generator.generate_all_formats(scan_id)
    else:
        print("[!] Invalid choice")
    
    print("\n[+] Report generation completed!")
